import docx
import pandas as pd

document = docx.Document('wordfile.docx')
para_text_list = []
for each_par in document.paragraphs:
    # print(each_par.text)
    para_text_list.append(each_par.text)

def create_df_from_table(c, tab):
    list_name = str(c)+"_result_list"
    list_name = []
    for i,each_row in enumerate(each_tab.rows):
        text = (each_cell.text for each_cell in each_row.cells)
        if ((c + 1) % 3) == 0:
            if i == 0:
                keys = tuple(text)
            else:
                each_dict_val = dict(zip(keys, text))
                list_name.append(each_dict_val)
        else:
            if i == 0: 
                pass
            elif i == 1:
                keys = tuple(text)
            else:
                each_dict_val = dict(zip(keys, text))
                list_name.append(each_dict_val)
    result_df = pd.DataFrame(list_name)
    return result_df

for c, each_tab in enumerate(document.tables):
    globals()[f'result_df_{c}'] = create_df_from_table(c, each_tab)
    
list_name = []
list_surname = []
list_ID = []
    
for i in range(0,len(para_text_list),1):
    text = para_text_list[i]
    key = "Name: "
    key_2 = "Surname: "
    key_3 = "ID: "
    
    try:
        name = text.index(key)
        txt = (text[name:])
        x = txt.replace("Name: ", "")
        list_name.append(str(x))
    except:
        pass
    
    try:
        surname = text.index(key_2)
        txt = (text[surname:])
        y = txt.replace("Surname: ", "")
        list_surname.append(str(y))
    except:
        pass
    
    try:
        ID = text.index(key_3)
        txt = (text[ID:])
        z = txt.replace("ID: ", "")
        list_ID.append(str(z))
    except:
        pass
    
k = 0
for j in range(0,len(list_name),1):
    form_info = [(j+1), list_name[j],list_surname[j],list_ID[j]]
    id_df = pd.DataFrame(form_info)
    with pd.ExcelWriter('docx2excel.xlsx', mode="a", engine="openpyxl", if_sheet_exists = "overlay") as writer: # First, you need create the "docx2excel.xlsx" then you can input data from word file.
        id_df.to_excel(writer, sheet_name = f'{j+1}. Form', startrow=4, startcol=2, index=False, header=False)
        globals()[f'result_df_{k}'].to_excel(writer, sheet_name = f'{j+1}. Form', startrow=13, startcol=1, index=False, header=False)
        globals()[f'result_df_{k+1}'].to_excel(writer, sheet_name = f'{j+1}. Form', startrow=13, startcol=6, index=False, header=False)
        globals()[f'result_df_{k+2}'].to_excel(writer, sheet_name = f'{j+1}. Form', startrow=22, startcol=1, index=False, header=False)
        k += 3 # k increases with 3 because one form of the wordfile.docx contains 3 tables. 
   
